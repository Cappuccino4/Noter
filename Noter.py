from tkinter import *
import sys, os
import os.path
from tkinter import filedialog
import glob
import re
from random import randint
from datetime import datetime
from docx import Document
from docx.shared import Pt
import calendar

# Colors!!
green = "#00c40a"
green_dark = "#008c07"
red = "#ff0000"
red_dark = "#a80000"
grey = "#424242"
grey_dark = "#2e2e2e"
dark_gray = "#1f1f1f"
main_color = "#635e5d"

top_color_list = ["#eb9e34", "#2bfff8", "#081cff", "#ff241c", "#07ad0f", "#0095ff", "#8c00ff", "#ff03dd"]
bottom_color_list = ["#b07627", "#18a19c", "#0714a6", "#99140f", "#057a0b", "#005694", "#5e00ab", "#960082"]

path = os.getenv("APPDATA")
path_NoterOptions = os.path.join(path, "NoterOptions")
path_Classes = os.path.join(path_NoterOptions, "Classes")


class FirstWindow():
    def __init__(self, master, title, size, color):
        # Defining variables
        self.master = master
        self.title = title
        self.size = size
        self.color = color
        self.master.title(self.title)
        self.master.geometry(self.size)
        self.master.config(background=self.color)
        self.master.resizable(width=False, height=False)

        # NOTER Text
        self.title_text = Text(self.master, font="fixedsys 50", width=5, height=1, padx=0,
                               pady=0, bg=main_color, bd=0)
        self.title_text.insert(INSERT, "NOTER")
        self.title_text.config(state=DISABLED)
        self.title_text.place(x=280, y=7)

        # Name Entry and text
        self.name_text = Text(self.master, font="fixedsys 30", height=1, padx=0,
                               pady=0, bg=main_color, bd=0)
        self.name_text.insert(INSERT, "Full Name:")
        self.name_text.config(state=DISABLED)
        self.name_text.place(x=70, y=100)

        self.name_entry = Entry(self.master, font="fixedsys 30", width=18, bd=0)
        self.name_entry.place(x=315, y=100)

        # Continue Button
        self.continue_button = Button(self.master, text="Continue", height=1, bg=green, command=self.buttonPressed,
                                      fg="white", font="fixedsys 50", activebackground=green_dark).place(x=200, y=300)

    def buttonPressed(self):
        if os.path.exists(path_NoterOptions):
            os.remove(path_NoterOptions + "/name.txt")
            f = open(path_NoterOptions + "/name.txt", "w+")
            f.write(self.name_entry.get())
            f.close()
        else:
            os.makedirs(path_NoterOptions)
            os.makedirs(path_Classes)
            f = open(path_NoterOptions + "/name.txt", "w+")
            f.write(self.name_entry.get())
            f.close()
        self.master.destroy()


class MainWindow():

    def __init__(self, master, title, size, color):
        # Defining variables
        self.master = master
        self.title = title
        self.size = size
        self.color = color
        self.master.title(self.title)
        self.master.geometry(self.size)
        self.master.config(background=self.color)
        self.master.resizable(width=False, height=False)
        self.match = ""
        self.updateList()

        # Top grey part
        self.top_panel = PanedWindow(master, bg=grey, height=93, width=800).place(x=0, y=0)

        # NOTER Text
        self.title_text = Text(self.top_panel, font="fixedsys 50", width=5, height=1, padx=0,
                               pady=0, bg=grey, bd=0)
        self.title_text.insert(INSERT, "NOTER")
        self.title_text.config(state=DISABLED)
        self.title_text.place(x=280, y=7)

        # Options button
        self.button_options = Button(self.top_panel, text="Options", bg=green, activebackground=green_dark, command=self.optionsButtonClicked,
                                    fg="white", font="fixedsys 17", activeforeground=grey).place(x=15, y=20)

        # New class button
        self.button_new_class = Button(self.top_panel, text="New Class", height=1, width=12, bg=green, command=self.classButtonClicked,
                                       fg="white", font="fixedsys 17", activebackground=green_dark, activeforeground=grey).place(x=590, y=18)

        # Randomize Colors
        self.random_colors_class = Button(self.master, text="New Colors!", command=self.relaunch, height=1, bg=red, fg="white",
                                  font="fixedsys 20", activebackground=red_dark, activeforeground=grey).place(x=15, y=530)

        # Quit Button
        self.button_exit = Button(self.master, text="Exit", command=self.on_cancel, height=3, width=8, bg=red, fg="white",
                                  font="fixedsys", activebackground=red_dark, activeforeground=grey).place(x=700, y=530)

    def classButtonClicked(self):
        obj = NewClassWindow(self, "New Class", "500x300")

    def optionsButtonClicked(self):
        obj = OptionMenu(self, "Options", "300x150")

    def updateList(self):

        count = -1
        x = 60
        y = 120
        last_button_width = 0
        self.class_name_list = []

        for file in glob.glob(path_Classes + "\*.txt"):
            count += 1
            x += last_button_width + 20
            temp_list = file.split("\\")
            class_name = temp_list[-1].replace(".txt", "")
            self.class_name_list.append(class_name)
            random_number = randint(0, 7)
            top_color = top_color_list[random_number]
            bottom_color = bottom_color_list[random_number]
            new_button = Button(self.master, text=class_name, height=1, bg=top_color, activebackground=bottom_color,
                       fg="white", font="fixedsys 22", command=lambda count=count: self.buttonPressed(count))
            new_button.place(x=x, y=y)
            last_button_width = new_button.winfo_reqwidth()
            if x + last_button_width > 700:
                x = 80
                y += 70
                new_button.place(x=x, y=y)

    def buttonPressed(self, c):
        class_name = self.class_name_list[c]

        current_date = datetime.now()

        month = current_date.strftime("%m")
        year = current_date.strftime("%Y")
        day = current_date.strftime("%d")

        # Create Notes Label
        notes_label = month + "_" + day + "_" + year

        date = day + " " + calendar.month_name[int(month)] + " " + year

        f = open(path_Classes + "/" + class_name + ".txt", "r")
        teacher_name = f.readline().rstrip()
        f.readline()
        path = f.readline()
        f.close()

        f = open(path_NoterOptions + "/name.txt", "r")
        name = f.readline().rstrip()

        finalpath = os.path.join(path, (notes_label + "_" + class_name + ".docx"))

        doc = Document()

        run1 = doc.add_paragraph().add_run(name)
        font = run1.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        run2 = doc.add_paragraph().add_run(teacher_name)
        font = run2.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        run3 = doc.add_paragraph().add_run(class_name)
        font = run3.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        run4 = doc.add_paragraph().add_run(date)
        font = run4.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        doc.save(finalpath)

        os.startfile(finalpath)
        self.master.destroy()
        exit()

    def relaunch(self):
        self.master.destroy()
        main()

    def on_cancel(self):
        self.master.destroy()


class NewClassWindow(Toplevel):

    def __init__(self, parent, title, size):
        super().__init__(name='new_class_window')
        self.parent = parent
        self.title(title)
        self.size = size
        self.filename = ""
        self.geometry(size)
        self.config(background=dark_gray)
        self.resizable(width=False, height=False)
        self.class_name = ""
        self.teacher_name = ""

        # Class Name Text and Entry
        self.class_name_text = Text(self, font="fixedsys 22", width=11, height=1, padx=0, pady=0, bd=0, fg="white", bg=dark_gray)
        self.class_name_text.insert(INSERT, "Class Name:")
        self.class_name_text.config(state=DISABLED)
        self.class_name_text.place(x=20, y=20)

        self.class_name_entry = Entry(self, font="fixedsys 22", width=18, bd=0)
        self.class_name_entry.place(x=200, y=20)

        # Teacher Name Text and Entry
        self.teacher_name_text = Text(self, font="fixedsys 22", height=1, padx=0, pady=0, bd=0, fg="white", bg=dark_gray)
        self.teacher_name_text.insert(INSERT, "Teacher's Name:")
        self.teacher_name_text.config(state=DISABLED)
        self.teacher_name_text.place(x=20, y=60)

        self.teacher_name_entry = Entry(self, font="fixedsys 22", width=14, bd=0)
        self.teacher_name_entry.place(x=264, y=60)


        # File Location
        self.directory_button = Button(self, text="Save Location", command=self.browseFiles, height=1, bg=green,
                       fg="white", font="fixedsys 30", activebackground=green_dark).place(x=80, y=120)

        # Continue Button
        self.continue_button = Button(self, text="Continue", height=1, bg=green, command=self.continueButtonClicked,
                       fg="white", font="fixedsys 22", activebackground=green_dark).place(x=170, y=210)

    def browseFiles(self):
        self.filename = filedialog.askdirectory(initialdir="/", title="Select a File")

    def continueButtonClicked(self):
        self.class_name = self.class_name_entry.get()
        self.teacher_name = self.teacher_name_entry.get()

        if self.class_name == "":
            self.destroy()
            MainWindow.on_cancel(self)
            main()
        else:
            f = open(path_Classes + "/" + self.class_name + ".txt", "w+")
            f.write(self.teacher_name + "\n" + self.class_name + "\n" + self.filename)
            f.close()
            self.destroy()
            MainWindow.on_cancel(self)
            main()


class OptionMenu(Toplevel):

    def __init__(self, parent, title, size):
        super().__init__(name='new_class_window')
        self.parent = parent
        self.title(title)
        self.size = size
        self.filename = ""
        self.geometry(size)
        self.config(background=dark_gray)
        self.resizable(width=False, height=False)

        # Delete class button
        self.delete_class_Button = Button(self, text="Delete Class", height=1, bg=green, command=self.deleteButtonPressed,
                       fg="white", font="fixedsys 22", activebackground=green_dark).place(x=45, y=15)

        # New Name Button
        self.change_name_button = Button(self, text="Change Name", height=1, bg=green, command=self.changeButtonPressed,
                       fg="white", font="fixedsys 22", activebackground=green_dark).place(x=52, y=80)

    def changeButtonPressed(self):
        self.destroy()
        MainWindow.on_cancel(self)
        firstTimeSetup = Tk()
        mainFenster = FirstWindow(firstTimeSetup, "Noter", "800x600", main_color)
        firstTimeSetup.mainloop()
        main()

    def deleteButtonPressed(self):
        file_name = self.browseFiles()
        os.remove(file_name)
        self.destroy()
        MainWindow.on_cancel(self)
        main()
        
    def browseFiles(self):
        self.filename = filedialog.askopenfilename(initialdir=path_Classes, title="Which Class do you want to delete?")
        return self.filename


def main():
    if os.path.exists(path_NoterOptions):
        mainWindow = Tk()
        mainFenster = MainWindow(mainWindow, "Noter", "800x600", main_color)
        mainWindow.mainloop()
    else:
        firstTimeSetup = Tk()
        mainFenster = FirstWindow(firstTimeSetup, "Noter", "800x600", main_color)
        firstTimeSetup.mainloop()
        main()

main()